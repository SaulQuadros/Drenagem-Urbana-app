#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import streamlit as st  
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def format_num(value, decimals=2):
    """
    Formata um número para exibir:
    - separador de milhar: ponto (.)
    - separador decimal: vírgula (,)
    """
    try:
        value = float(value)
        # Exemplo: 1234567.89 -> "1,234,567.89"
        formatted = f"{value:,.{decimals}f}"
        # Inverte os separadores: vírgula para ponto e ponto para vírgula
        formatted = formatted.replace(",", "X").replace(".", ",").replace("X", ".")
        return formatted
    except Exception:
        return value

# Inicializa o estado para as variáveis, se não existirem
if "tc" not in st.session_state:
    st.session_state.tc = None
if "i_max" not in st.session_state:
    st.session_state.i_max = None
if "Q" not in st.session_state:
    st.session_state.Q = None
if "P_n_percent" not in st.session_state:
    st.session_state.P_n_percent = None

# Inicializa os campos dos Dados do Projeto, se não existirem
if "nome_projeto" not in st.session_state:
    st.session_state.nome_projeto = ""
if "tecnico" not in st.session_state:
    st.session_state.tecnico = ""
if "resumo" not in st.session_state:
    st.session_state.resumo = ""

# (Opcional) Inicializa também outros campos que serão usados em Cálculos
if "area_km2_bacia" not in st.session_state:
    st.session_state.area_km2_bacia = 4.5
if "perimetro_km" not in st.session_state:
    st.session_state.perimetro_km = 9.6
if "comprimento_curso_principal_km" not in st.session_state:
    st.session_state.comprimento_curso_principal_km = 3.2
if "comprimento_retalinea_km" not in st.session_state:
    st.session_state.comprimento_retalinea_km = 2.5
if "comprimento_total_cursos_agua_km" not in st.session_state:
    st.session_state.comprimento_total_cursos_agua_km = 9.0
if "desnivel_m" not in st.session_state:
    st.session_state.desnivel_m = 25.0

# Título na barra lateral
st.sidebar.title("Drenagem Urbana")

# Menu principal utilizando selectbox (sem definir index fixo)
opcao_principal = st.sidebar.selectbox(
    "Selecione a Opção",
    ["Dados do Projeto", "Cálculos"],
    key="menu_principal"
)

# --- DADOS DO PROJETO ---
if opcao_principal == "Dados do Projeto":
    st.title("Dados do Projeto")
    
    # Os widgets usam os valores armazenados em st.session_state
    st.text_input("Nome do Projeto", max_chars=100, key="nome_projeto")
    st.text_input("Técnico Responsável", max_chars=100, key="tecnico")
    st.text_area("Resumo", max_chars=200, height=90, key="resumo")
    
# --- CÁLCULOS ---
elif opcao_principal == "Cálculos":
    # Submenu com os tipos de cálculos disponíveis
    menu = st.sidebar.radio(
        "Selecione o tipo de Cálculo", 
        ["Característica da Bacia", "Microdrenagem - Método Racional"],
        key="submenu_calculos"
    )
    
    # --- Relatório de Parâmetros da Bacia ---
    if menu == "Característica da Bacia":
        st.title('Parâmetros de Bacia Hidrográfica')
        
        st.sidebar.header('Insira os dados da bacia')
        area_km2 = st.sidebar.number_input(
            'Área da Bacia (km²)', 
            min_value=0.01, value=4.5, step=0.01, format="%.2f", 
            key="area_km2_bacia"
        )
        perimetro_km = st.sidebar.number_input(
            'Perímetro da Bacia (km)', 
            min_value=0.1, value=9.6, step=0.1, format="%.2f", 
            key="perimetro_km"
        )
        comprimento_curso_principal_km = st.sidebar.number_input(
            'Comprimento do Curso Principal (km)', 
            min_value=0.1, value=3.2, step=0.1, format="%.2f", 
            key="comprimento_curso_principal_km"
        )
        comprimento_retalinea_km = st.sidebar.number_input(
            'Comprimento em Linha Reta (km)', 
            min_value=0.1, value=2.5, step=0.1, format="%.2f", 
            key="comprimento_retalinea_km"
        )
        comprimento_total_cursos_agua_km = st.sidebar.number_input(
            "Comprimento Total dos Cursos d'Água (km)", 
            min_value=1.0, value=9.0, step=0.1, format="%.2f", 
            key="comprimento_total_cursos_agua_km"
        )
        desnivel_m = st.sidebar.number_input(
            'Desnível da Bacia (m)', 
            min_value=1.0, value=25.0, step=1.0, format="%.2f", 
            key="desnivel_m"
        )
        
        # Cálculos dos parâmetros
        kf = area_km2 / (comprimento_curso_principal_km ** 2)
        kc = 0.28 * perimetro_km / (area_km2 ** 0.5)
        dd = comprimento_total_cursos_agua_km / area_km2
        lm = area_km2 / (4 * comprimento_total_cursos_agua_km)
        sc = comprimento_curso_principal_km / comprimento_retalinea_km
        dc = (desnivel_m / (comprimento_curso_principal_km * 1000)) * 100
        
        resultados = [
            (
                "Coeficiente de Forma (Kf)",
                kf,
                "quanto mais próximo de 1, mais arredondada é a bacia, indicando picos de vazões mais elevados e maior tendência para enchentes rápidas, sendo o oposto para valores que se aproximam de 0."
            ),
            (
                "Coeficiente de Compacidade (Kc)",
                kc,
                "quanto mais próximo de 1, mais circular é o formato da bacia e favorece o escoamento com altos picos de vazão, sendo a bacia mais sujeita a inundações rápidas, sendo o oposto para valores que se afastam de 1."
            ),
            (
                "Densidade de Drenagem (Dd)",
                dd,
                "valores maiores que 1 indicam maior rapidez no escoamento superficial e menor infiltração, com maior risco de enchentes, e o inverso para valores menores que 1."
            ),
            (
                "Extensão Média do Escoamento (lm)",
                lm,
                "valores entre 100m e 250m indicam uma bacia com drenagem moderada, com equilíbrio entre infiltração e escoamento superficial, contudo, abaixo de 100 m, o escoamento superficial tende a ser rápido com pico de vazões elevados, e acima de 250 m o inverso."
            ),
            (
                "Índice de Sinuosidade (Sc)",
                sc,
                "valores próximos de 1 indicam canais mais retos e maior eficiência de drenagem, portanto, quanto maior o valor maior a sinuosidade e com isso, maior risco de enchentes."
            ),
            (
                "Declividade do Curso d'água Principal (Dc)",
                dc,
                "valores abaixo de 1% indicam maior risco de enchentes, pois a drenagem é demorada, sendo rios de planícies, e acima de 5% indicam rios com corredeiras e elevada velocidade de escoamento."
            )
        ]
        
        st.header('Resultados dos Parâmetros da Bacia')
        st.markdown(f'''
        - **Coeficiente de Forma (Kf)**: {format_num(kf, 3)}  
          **Interpretação**: quanto mais próximo de 1, mais arredondada é a bacia, indicando picos de vazões mais elevados e maior 
          tendência para enchentes rápidas, sendo o oposto para valores que se aproximam de 0.
        
        - **Coeficiente de Compacidade (Kc)**: {format_num(kc, 3)}  
          **Interpretação**: quanto mais próximo de 1, mais circular é o formato da bacia e favorece o escoamento com altos picos de vazão, 
          sendo a bacia mais sujeita a inundações rápidas, sendo o oposto para valores que se afastam de 1.
        
        - **Densidade de Drenagem (Dd)**: {format_num(dd, 3)} km/km²  
          **Interpretação**: valores maiores que 1 indicam maior rapidez no escoamento superficial e menor infiltração, com maior risco de 
          enchentes, e o inverso para valores menores que 1.
        
        - **Extensão Média do Escoamento (lm)**: {format_num(lm, 3)} km
          **Interpretação**: valores entre 100m e 250m indicam uma bacia com drenagem moderada, com equilíbrio entre infiltração e escoamento 
          superficial, contudo, abaixo de 100 m, o escoamento superficial tende a ser rápido com pico de vazões elevados, e acima de 250 m 
          o inverso.
        
        - **Índice de Sinuosidade (Sc)**: {format_num(sc, 3)}  
          **Interpretação**: valores próximos de 1 indicam canais mais retos e maior eficiência de drenagem, portanto, quanto maior o valor 
          maior a sinuosidade e com isso, maior risco de enchentes.
        
        - **Declividade do Curso d'água Principal (Dc)**: {format_num(dc, 3)}%  
          **Interpretação**: valores abaixo de 1% indicam maior risco de enchentes, pois a drenagem é demorada, sendo rios de planícies, 
          e acima de 5% indicam rios com corredeiras e elevada velocidade de escoamento. 
        ''')
        
        # Geração do documento Word para Parâmetros da Bacia com os Dados do Projeto no início
        if st.button('📄 Gerar Relatório Word - Parâmetros da Bacia', key="bt_rel_bacia"):
            doc = Document()
        
            sec = doc.sections[0]
            sec.top_margin = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin = Cm(2.5)
            sec.right_margin = Cm(2.5)
        
            # Inserindo os Dados do Projeto
            doc.add_heading('Dados do Projeto', level=1)
            doc.add_paragraph(f"Nome do Projeto: {st.session_state.get('nome_projeto', 'Não informado')}")
            doc.add_paragraph(f"Técnico Responsável: {st.session_state.get('tecnico', 'Não informado')}")
            doc.add_paragraph(f"Resumo: {st.session_state.get('resumo', 'Não informado')}")
            doc.add_paragraph()
        
            titulo = doc.add_heading('Relatório de Parâmetros da Bacia Hidrográfica', 0)
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            titulo.runs[0].font.size = Pt(16)
            titulo.runs[0].bold = True
            titulo.runs[0].font.name = 'Aptos'
        
            doc.add_paragraph()
        
            for nome, valor, interpretacao in resultados:
                p_param = doc.add_paragraph()
                run_param = p_param.add_run(f"{nome}: ")
                run_param.bold = True
                run_param.font.size = Pt(11)
                run_param.font.name = 'Aptos'
                run_valor = p_param.add_run(f"{format_num(valor, 3)}")
                run_valor.font.size = Pt(11)
                run_valor.font.name = 'Aptos'
                p_param.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p_param.paragraph_format.space_after = Pt(6)
                
                p_interp = doc.add_paragraph()
                run_interp_label = p_interp.add_run("Interpretação: ")
                run_interp_label.bold = True
                run_interp_label.font.size = Pt(11)
                run_interp_label.font.name = 'Aptos'
                run_interp_text = p_interp.add_run(interpretacao)
                run_interp_text.font.size = Pt(11)
                run_interp_text.font.name = 'Aptos'
                p_interp.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p_interp.paragraph_format.space_after = Pt(12)
        
            doc.save("relatorio_bacia.docx")
        
            with open("relatorio_bacia.docx", "rb") as f:
                st.download_button("⬇️ Baixar relatório", f, file_name="relatorio_bacia.docx")
    
    # --- Relatório de Microdrenagem - Método Racional ---
    elif menu == "Microdrenagem - Método Racional":
        st.title("Microdrenagem - Método Racional")
        
        st.markdown("### Escolha do Modelo de Tempo de Concentração")
        modelo_tc = st.selectbox(
            "Selecione o modelo para o cálculo do tempo de concentração:",
            ["Kirpich", "Kirpich Modificado", "Van Te Chow", "George Ribeiro", "Piking", "USACE", "DNOS", "NRCS (SCS)"],
            key="modelo_tc"
        )
        
        # Inputs para os modelos – L em km e H em m
        if modelo_tc == "Kirpich":
            st.markdown("#### Parâmetros para a fórmula de Kirpich")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H")
            st.session_state.tc = 57 * (((L_km ** 3) / H) ** 0.385)
        elif modelo_tc == "Kirpich Modificado":
            st.markdown("#### Parâmetros para a fórmula de Kirpich Modificado")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km_mod")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H_mod")
            st.session_state.tc = 85.2 * (((L_km ** 3) / H) ** 0.385)
        elif modelo_tc == "Van Te Chow":
            st.markdown("#### Parâmetros para a fórmula de Van Te Chow")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km_vtc")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H_vtc")
            S = H / (L_km * 1000)
            st.session_state.tc = 5.773 * ((L_km / (S ** 0.5)) ** 0.64)
        elif modelo_tc == "George Ribeiro":
            st.markdown("#### Parâmetros para a fórmula de George Ribeiro")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km_gr")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H_gr")
            S = H / (L_km * 1000)
            pr = st.number_input("Parâmetro (pr) - Porção da bacia coberta por vegetação", min_value=0.0, max_value=1.0, value=0.5, step=0.01, key="pr")
            st.session_state.tc = (16 * L_km) / ((1.05 - 0.2 * pr) * ((100 * S) ** 0.04))
        elif modelo_tc == "Piking":
            st.markdown("#### Parâmetros para a fórmula de Piking")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km_piking")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H_piking")
            S = H / (L_km * 1000)
            st.session_state.tc = 5.3 * (((L_km ** 2) / S) ** (1/3))
        elif modelo_tc == "USACE":
            st.markdown("#### Parâmetros para a fórmula de USACE")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km_usace")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H_usace")
            S = H / (L_km * 1000)
            st.session_state.tc = 7.504 * (L_km ** 0.76) * (S ** (-0.19))
        elif modelo_tc == "DNOS":
            st.markdown("#### Parâmetros para a fórmula de DNOS")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km_dnos")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H_dnos")
            S = H / (L_km * 1000)
            A = st.session_state.get("area_km2_micro", 1.0)
            terreno_options = [
                "arenoso-argiloso, coberto de vegetação intensa, elevada absorção",
                "comum, coberto de vegetação, absorção apreciável",
                "argiloso, coberto de vegetação, absorção média",
                "argiloso de vegetação média, pouca absorção",
                "com rocha, escassa vegetação, baixa absorção",
                "Rochoso, vegetação rala, reduzida absorção"
            ]
            terreno = st.selectbox("Selecione o tipo de terreno", terreno_options, key="terreno")
            if terreno == terreno_options[0]:
                K = 2.0
            elif terreno == terreno_options[1]:
                K = 3.0
            elif terreno == terreno_options[2]:
                K = 4.0
            elif terreno == terreno_options[3]:
                K = 4.5
            elif terreno == terreno_options[4]:
                K = 5.0
            elif terreno == terreno_options[5]:
                K = 5.5
            st.session_state.tc = (10 / K) * (((100 * A ** 0.3) * (L_km ** 0.2)) / (S ** 0.4))
        elif modelo_tc == "NRCS (SCS)":
            st.markdown("#### Parâmetros para a fórmula de NRCS (SCS)")
            L_km = st.number_input("Comprimento máximo do percurso d'água (km)", min_value=0.1, value=1.0, step=0.1, key="L_km_nrcs")
            H = st.number_input("Desnível da bacia (m)", min_value=1.0, value=20.0, step=1.0, key="H_nrcs")
            S = H / (L_km * 1000)
            area_tipo = st.selectbox("Tipo de Área", ["Urbana", "Rural"], key="area_tipo")
            cond_area = st.selectbox("Condição da Área", ["Seco", "Úmido"], key="cond_area")
            if area_tipo == "Urbana":
                uso = st.selectbox("Uso do Solo", ["100% pavimentadas", "Urbanas altamente impermeáveis", "Residenciais", "Com parques"], key="uso_urbano")
                if uso == "100% pavimentadas":
                    CN = 98 if cond_area=="Seco" else 99
                elif uso == "Urbanas altamente impermeáveis":
                    CN = 85 if cond_area=="Seco" else 95
                elif uso == "Residenciais":
                    CN = 70 if cond_area=="Seco" else 85
                elif uso == "Com parques":
                    CN = 60 if cond_area=="Seco" else 75
            else:
                uso = st.selectbox("Uso do Solo", ["Pastagem", "Solo argiloso", "Florestas densas", "Solo compactado"], key="uso_rural")
                if uso == "Pastagem":
                    CN = 39 if cond_area=="Seco" else 61
                elif uso == "Solo argiloso":
                    CN = 66 if cond_area=="Seco" else 85
                elif uso == "Florestas densas":
                    CN = 30 if cond_area=="Seco" else 55
                elif uso == "Solo compactado":
                    CN = 75 if cond_area=="Seco" else 85
            st.session_state.tc = 3.42 * ((1000 / CN - 9) ** 0.7) * (L_km ** 0.8) * (S ** (-0.5))
        else:
            st.info("Selecione um modelo válido.")
            st.session_state.tc = None
        
        st.markdown("### Dados para o Cálculo da Intensidade Pluviométrica Máxima")
        a = st.number_input("Coeficiente a", value=1000.0, step=10.0, key="a")
        b = st.number_input("Coeficiente b", value=10.0, step=0.01, key="b")
        m = st.number_input("Expoente m", value=1.0, step=0.01, key="m")
        n = st.number_input("Expoente n", value=1.0, step=0.01, key="n")
        
        # Novos inputs para a equação de i_max e probabilidade
        T = st.number_input("Tempo de Retorno (anos)", min_value=1, max_value=1000, value=10, step=1, key="T")
        n_period = st.number_input("Período de análise (n anos)", min_value=1, max_value=T, value=1, step=1, key="n_period")
        
        st.markdown("### Coeficiente de Escoamento Superficial (C)")
        C = st.number_input("Insira o valor de C", value=0.7, step=0.01, key="C")
        
        st.markdown("### Dados da Bacia para o Método Racional")
        area_km2_md = st.number_input("Área da Bacia (km²)", min_value=0.001, value=1.0, step=0.001, key="area_km2_micro")
        area_m2 = area_km2_md * 1e6
        
        # Botão de cálculo
        if st.button("Calcular", key="calcular"):
            if st.session_state.tc is None:
                st.error("Selecione um modelo de tempo de concentração implementado.")
            else:
                td = st.session_state.tc  # Considera td = tc
                try:
                    st.session_state.i_max = (a * (T ** m)) / ((td + b) ** n)
                except Exception as e:
                    st.error("Erro no cálculo da intensidade: verifique os valores inseridos.")
                    st.session_state.i_max = None
                
                if st.session_state.i_max is not None:
                    P = 1 / T
                    P_n = 1 - ((1 - P) ** n_period)
                    st.session_state.P_n_percent = P_n * 100
                    
                    i_max_ms = st.session_state.i_max * 2.78e-7
                    st.session_state.Q = C * i_max_ms * area_m2
                    
                    st.markdown("#### Resultados do Projeto")
                    st.write(f"Tempo de Concentração (tc = td): **{format_num(td, 2)} minutos**")
                    st.write(f"Intensidade Pluviométrica Máxima (i_max): **{format_num(st.session_state.i_max, 2)} mm/h**")
                    st.write(f"Vazão Máxima de Projeto (Q): **{format_num(st.session_state.Q, 3)} m³/s**")
                    st.write(f"Probabilidade de ocorrência em {n_period} ano(s): **{format_num(st.session_state.P_n_percent, 2)}%**")
        
        if st.button("📄 Gerar Relatório Word - Microdrenagem", key="bt_rel_micro"):
            if (st.session_state.tc is None or
                st.session_state.i_max is None or
                st.session_state.Q is None or
                st.session_state.P_n_percent is None):
                st.error("Realize o cálculo primeiro para gerar o relatório.")
            else:
                # Variáveis auxiliares para os valores de L_km e H
                L_km_val = (st.session_state.get('L_km') or 
                            st.session_state.get('L_km_mod') or 
                            st.session_state.get('L_km_vtc') or 
                            st.session_state.get('L_km_gr') or 
                            st.session_state.get('L_km_piking') or 
                            st.session_state.get('L_km_usace') or 
                            st.session_state.get('L_km_dnos') or 
                            st.session_state.get('L_km_nrcs') or '')
                H_val = (st.session_state.get('H') or 
                         st.session_state.get('H_mod') or 
                         st.session_state.get('H_vtc') or 
                         st.session_state.get('H_gr') or 
                         st.session_state.get('H_piking') or 
                         st.session_state.get('H_usace') or 
                         st.session_state.get('H_dnos') or 
                         st.session_state.get('H_nrcs') or '')
                
                doc = Document()
                sec = doc.sections[0]
                sec.top_margin = Cm(2.0)
                sec.bottom_margin = Cm(2.0)
                sec.left_margin = Cm(2.5)
                sec.right_margin = Cm(2.5)
        
                # Inserindo os Dados do Projeto no início do documento
                doc.add_heading('Dados do Projeto', level=1)
                doc.add_paragraph(f"Nome do Projeto: {st.session_state.get('nome_projeto', 'Não informado')}")
                doc.add_paragraph(f"Técnico Responsável: {st.session_state.get('tecnico', 'Não informado')}")
                doc.add_paragraph(f"Resumo: {st.session_state.get('resumo', 'Não informado')}")
                doc.add_paragraph()
        
                titulo = doc.add_heading('Microdrenagem - Método Racional', 0)
                titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                titulo.runs[0].font.size = Pt(16)
                titulo.runs[0].bold = True
                titulo.runs[0].font.name = 'Aptos'
        
                doc.add_paragraph()
        
                # Seção: Dados do Projeto (relatório interno)
                doc.add_heading('Dados do Projeto', level=2)
                dados_projeto = [
                    f"Modelo de Cálculo do tc: {modelo_tc}",
                    f"Comprimento máximo do percurso d'água (km): {L_km_val}",
                    f"Desnível da bacia (m): {H_val}",
                    f"Tempo de Concentração (tc = td): {format_num(st.session_state.tc, 2)} minutos",
                    f"Coeficiente a: {st.session_state.get('a', '')}",
                    f"Coeficiente b: {st.session_state.get('b', '')}",
                    f"Expoente m: {st.session_state.get('m', '')}",
                    f"Expoente n: {st.session_state.get('n', '')}",
                    f"Tempo de Retorno (T): {st.session_state.get('T', '')} ano(s)",
                    f"Período de análise (n anos): {st.session_state.get('n_period', '')}",
                    f"Coeficiente de Escoamento (C): {st.session_state.get('C', '')}",
                    f"Área da Bacia (km²): {st.session_state.get('area_km2_micro', '')}"
                ]
                for item in dados_projeto:
                    doc.add_paragraph(item, style='List Bullet')
        
                doc.add_paragraph()  # Espaço entre seções
        
                # Seção: Resultados
                doc.add_heading('Resultados', level=2)
                resultados_rel = [
                    f"Tempo de Concentração (tc = td): {format_num(st.session_state.tc, 2)} minutos",
                    f"Intensidade Pluviométrica Máxima (i_max): {format_num(st.session_state.i_max, 2)} mm/h",
                    f"Vazão Máxima de Projeto (Q): {format_num(st.session_state.Q, 3)} m³/s",
                    f"Probabilidade de ocorrência em {st.session_state.get('n_period', '')} ano(s): {format_num(st.session_state.P_n_percent, 2)}%"
                ]
                for item in resultados_rel:
                    doc.add_paragraph(item, style='List Bullet')
        
                doc.save("relatorio_vazao_maxima.docx")
        
                with open("relatorio_vazao_maxima.docx", "rb") as f:
                    st.download_button("⬇️ Baixar relatório", f, file_name="relatorio_vazao_maxima.docx")
                
                st.markdown("#### Resultados do Projeto (mantidos na tela)")
                st.write(f"Tempo de Concentração (tc = td): **{format_num(st.session_state.tc, 2)} minutos**")
                st.write(f"Intensidade Pluviométrica Máxima (i_max): **{format_num(st.session_state.i_max, 2)} mm/h**")
                st.write(f"Vazão Máxima de Projeto (Q): **{format_num(st.session_state.Q, 3)} m³/s**")
                st.write(f"Probabilidade de ocorrência em {st.session_state.get('n_period', '')} ano(s): **{format_num(st.session_state.P_n_percent, 2)}%**")

